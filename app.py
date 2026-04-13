import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import yaml
from datetime import datetime
import sqlite3
import json
from pathlib import Path
import logging
import traceback
from contextlib import contextmanager
from functools import wraps
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, KeepTogether, PageBreak
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 自定义异常类，用于区分不同类型的应用错误
class AppError(Exception):
    """应用程序基础异常类"""
    pass


class DatabaseError(AppError):
    """数据库操作相关异常"""
    pass


class ConfigError(AppError):
    """配置加载和处理相关异常"""
    pass


class FileError(AppError):
    """文件操作相关异常"""
    pass


class BusinessLogicError(AppError):
    """业务逻辑相关异常"""
    pass


class ExportError(AppError):
    """报告导出相关异常"""
    pass

# 异常处理装饰器
def handle_errors(default_return_value=None, error_type=AppError):
    """统一异常处理装饰器
    
    封装重复的错误处理逻辑，记录详细日志并转换为自定义异常。
    
    参数:
        default_return_value: 异常发生时的默认返回值
        error_type: 转换后的自定义异常类型
    
    用法示例:
        @handle_errors(default_return_value=0)
        def calculate_score(data):
            # 业务逻辑
            return result
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                # 记录详细错误日志
                logging.error(f"函数 {func.__name__} 执行失败: {str(e)}")
                logging.error(traceback.format_exc())
                # 如果指定了默认返回值，则返回
                if default_return_value is not None:
                    # 支持可调用对象作为默认返回值
                    if callable(default_return_value):
                        return default_return_value()
                    return default_return_value
                # 否则抛出自定义异常
                raise error_type(str(e)) from e
        return wrapper
    return decorator

# 新增：导入统一配置加载模块
from config_loader import config

# 全局章节名称映射
SECTION_MAPPING = {
    'organization_context': '组织环境',
    'leadership': '领导作用',
    'planning': '策划',
    'support': '支持',
    'operation': '运行',
    'performance_evaluation': '绩效评价',
    'improvement': '改进'
}

# 语言辅助函数
@handle_errors(default_return_value=None, error_type=ConfigError)
def get_lang_text(key):
    """获取当前语言环境下的文本
    
    参数:
        key: 文本键名
    
    返回:
        str: 当前语言环境下的文本
    """
    if getattr(st.session_state, 'language', 'zh') == 'zh':
        return config.lang_zh[key]
    return config.lang_en[key]

def get_language():
    """获取当前语言环境
    
    返回:
        str: 当前语言环境 ('zh' 或 'en')
    """
    return getattr(st.session_state, 'language', 'zh')

@handle_errors(default_return_value=lambda: get_lang_text('unknown_section'), error_type=ConfigError)
def get_section_title(section_id, language=None):
    """获取章节标题
    
    参数:
        section_id: 章节ID
        language: 语言环境 (可选)
    
    返回:
        str: 章节标题
    """
    lang = language or get_language()
    if lang == 'en':
        return section_id.replace('_', ' ').title()
    return SECTION_MAPPING.get(section_id, section_id)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# 定义数据库文件名
DATABASE_NAME = 'assessment_data.db'

# 数据库上下文管理器
@contextmanager
@handle_errors(error_type=DatabaseError)
def get_db_connection(db_name=None):
    """数据库连接上下文管理器
    
    提供安全的数据库连接管理，确保连接在使用后正确关闭。
    自动处理连接错误和异常情况，并记录详细日志。
    
    参数:
        db_name (str, optional): 数据库文件名，默认使用DATABASE_NAME
    
    用法示例:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute("SELECT * FROM assessment_results")
            results = c.fetchall()
    """
    database = db_name or DATABASE_NAME
    conn = sqlite3.connect(database)
    try:
        yield conn
    finally:
        if conn:
            conn.close()

# 初始化数据库
@handle_errors(error_type=DatabaseError)
def init_db():
    """初始化数据库"""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS assessment_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                responses TEXT,
                sub_responses TEXT
            )
        ''')
        conn.commit()
    logging.info("数据库初始化成功")

# 保存评估结果
@handle_errors(error_type=DatabaseError)
def save_assessment_results(responses, sub_responses):
    """保存评估结果到数据库
    
    将评估结果（主回答和子回答）保存到SQLite数据库中，并记录时间戳。
    
    参数:
        responses (dict): 主问题回答字典，键为问题ID，值为回答内容
        sub_responses (dict): 子问题回答字典，键为子问题ID，值为回答内容
    
    返回:
        None
    
    异常:
        DatabaseError: 当保存失败时抛出
        
    用法示例:
        test_responses = {'org_1': 'yes', 'org_2': 'no', 'lead_1': '5'}
        test_sub_responses = {'org_1_1': 'yes', 'org_1_2': 'no'}
        save_assessment_results(test_responses, test_sub_responses)
    """
    logging.info("=== 开始保存评估结果 ===")
    logging.debug("连接数据库...")
    with get_db_connection() as conn:
        logging.info("✓ 数据库连接成功")
        c = conn.cursor()
        logging.debug("准备执行SQL插入...")
        
        # 执行插入并获取受影响的行数
        c.execute('''
            INSERT INTO assessment_results (timestamp, responses, sub_responses)
            VALUES (?, ?, ?)
        ''', (datetime.now().isoformat(), 
              json.dumps(responses), 
              json.dumps(sub_responses)))
        
        affected_rows = conn.total_changes
        logging.info("✓ SQL插入执行完成")
        logging.info(f"受影响的行数: {affected_rows}")
        
        conn.commit()
        logging.info("✓ 事务提交成功")
        
        # 在同一连接中验证数据是否保存成功
        logging.debug("=== 在同一连接中验证数据 ===")
        c.execute("SELECT COUNT(*) FROM assessment_results")
        count_in_same_conn = c.fetchone()[0]
        logging.debug(f"同一连接中的记录数: {count_in_same_conn}")
        
        c.execute("SELECT * FROM assessment_results ORDER BY timestamp DESC LIMIT 1")
        result_in_same_conn = c.fetchone()
        logging.debug(f"同一连接中的查询结果: {result_in_same_conn}")
    logging.info("结果保存成功")
    logging.info("=== 评估结果保存完成 ===")

# 加载最近的评估结果
@handle_errors(default_return_value=({}, {}), error_type=DatabaseError)
def load_latest_assessment_results():
    """加载最近的评估结果
    
    从数据库中加载最新保存的评估结果（按时间戳降序排序）。
    
    参数:
        None
    
    返回:
        tuple: 包含两个字典的元组
            - responses (dict): 主问题回答字典，键为问题ID，值为回答内容
            - sub_responses (dict): 子问题回答字典，键为子问题ID，值为回答内容
    
    异常处理:
        - 如果数据库连接失败或查询出错，返回空字典对 ({}, {})
        - 如果数据库中没有记录，返回空字典对 ({}, {})
    
    用法示例:
        responses, sub_responses = load_latest_assessment_results()
        if responses:
            print("加载到最近的评估结果")
        else:
            print("没有找到评估结果")
    """
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('''
            SELECT responses, sub_responses 
            FROM assessment_results 
            ORDER BY timestamp DESC 
            LIMIT 1
        ''')
        result = c.fetchone()
        
        if result:
            return json.loads(result[0]), json.loads(result[1])
        return {}, {}

# 初始化会话状态
@handle_errors(default_return_value=None, error_type=BusinessLogicError)
def init_session_state():
    """初始化Streamlit会话状态
    
    设置评估工具所需的所有会话状态变量及其默认值。
    
    会话状态变量包括：
    - responses: 主问题回答字典，默认空字典
    - sub_responses: 子问题回答字典，默认空字典
    - last_save_time: 最后保存时间，默认当前时间
    - force_refresh: 强制刷新标志，默认False
    - language: 显示语言，默认'zh'（中文）
    
    参数:
        None
    
    返回:
        None
    
    用法示例:
        init_session_state()  # 在应用启动时调用初始化所有会话状态
    """
    if 'responses' not in st.session_state:
        st.session_state.responses = {}
    if 'sub_responses' not in st.session_state:
        st.session_state.sub_responses = {}
    if 'last_save_time' not in st.session_state:
        st.session_state.last_save_time = datetime.now()
    if 'force_refresh' not in st.session_state:
        st.session_state.force_refresh = False
    if 'language' not in st.session_state:
        st.session_state.language = 'zh'  # 默认中文

# 加载评估问题
@handle_errors(error_type=ConfigError)
def load_questionnaire():
    """加载评估问题
    
    从配置文件加载中英文评估问题，并合并为统一的格式供前端使用。
    支持ISO 55001标准的所有章节问题。
    
    返回:
        dict: 合并后的评估问题字典，结构如下：
            {
                '章节英文键名': {
                    'id': '章节ID',
                    'name': {'zh': '中文章节名', 'en': '英文章节名'},
                    'questions': {
                        '问题ID': {
                            'type': '问题类型(XO/PJ/PW)',
                            'description': {'zh': '中文描述', 'en': '英文描述'},
                            'sub_questions': {'zh': ['中文子问题列表'], 'en': ['英文子问题列表']}  # PW类型特有
                        }
                    }
                }
            }
            
    加载逻辑:
        1. 从config/questionnaire.yaml加载中文问题
        2. 从config/questionnaire_en.yaml加载英文问题
        3. 创建章节名称映射关系
        4. 合并中英文内容，保留统一结构
        5. 记录加载日志
        
    示例:
        questionnaire = load_questionnaire()
        # 获取所有章节
        sections = list(questionnaire.keys())
        # 获取第一个章节的第一个问题
        first_section = list(questionnaire.values())[0]
        first_question = list(first_section['questions'].values())[0]
    """
    # 加载中文问题配置文件
    with open('config/questionnaire.yaml', 'r', encoding='utf-8') as file:
        questions_zh = yaml.safe_load(file)
    
    # 加载英文问题配置文件
    with open('config/questionnaire_en.yaml', 'r', encoding='utf-8') as file:
        questions_en = yaml.safe_load(file)
    
    logging.info("成功加载评估问题")
    
    # 合并中英文内容为统一格式
    formatted_questions = {}
    
    # 使用全局章节名称映射
    section_mapping = SECTION_MAPPING
    
    # 遍历英文问题作为基准
    for section_en, section_data_en in questions_en.items():
        section_zh = section_mapping.get(section_en)
        if not section_zh:
            logging.warning(f"找不到章节 '{section_en}' 的中文映射")
            continue
            
        section_zh_data = questions_zh.get(section_zh, {})
        
        formatted_questions[section_en] = {
            'id': section_en,
            'name': {
                'zh': section_zh,
                'en': section_en.replace('_', ' ').title()
            },
            'questions': {}
        }
        
        for q_id, q_data_en in section_data_en.items():
            q_data_zh = section_zh_data.get(q_id, {})
            
            formatted_question = {
                'type': q_data_en['type'],
                'description': {
                    'en': q_data_en['description'],
                    'zh': q_data_zh.get('description', q_data_en['description'])  # 如果没有中文描述，使用英文
                }
            }
            
            # 处理子问题
            if 'sub_questions' in q_data_en:
                formatted_question['sub_questions'] = {
                    'en': q_data_en['sub_questions'],
                    'zh': q_data_zh.get('sub_questions', q_data_en['sub_questions'])  # 如果没有中文子问题，使用英文
                }
            
            formatted_questions[section_en]['questions'][q_id] = formatted_question
    
    return formatted_questions

# 加载预设分值
@handle_errors(default_return_value=None, error_type=ConfigError)
def load_preset_scores():
    """加载预设分值"""
    with open('config/preset_scores.yaml', 'r', encoding='utf-8') as file:
        preset_scores = yaml.safe_load(file)
    logging.info("成功加载预设分值")
    return preset_scores

# 应用预设分值
@handle_errors(error_type=BusinessLogicError)
def apply_preset_scores(grade):
    """应用预设分值到会话状态"""
    preset_scores = load_preset_scores()
    if not preset_scores:
        st.error(get_lang_text('cannot_load_preset_scores'))
        return
    
    # 清空当前会话状态
    st.session_state.responses = {}
    st.session_state.sub_responses = {}
    
    # 应用主问题分值
    if 'preset_scores' in preset_scores and grade in preset_scores['preset_scores']:
        for section_en, questions in preset_scores['preset_scores'][grade].items():
            for q_id, value in questions.items():
                key = f"{section_en}_{q_id}"
                if isinstance(value, str):  # XO类型
                    st.session_state.responses[key] = 4 if value == 'yes' else 0
                else:  # PJ类型
                    st.session_state.responses[key] = value
    
    # 应用子问题分值
    if 'preset_sub_scores' in preset_scores and grade in preset_scores['preset_sub_scores']:
        for section_en, questions in preset_scores['preset_sub_scores'][grade].items():
            for q_id, sub_values in questions.items():
                if isinstance(sub_values, list):
                    key = f"{section_en}_{q_id}"
                    for i, selected in enumerate(sub_values, 1):
                        sub_key = f"{key}_sub_{i}"
                        st.session_state.sub_responses[sub_key] = selected
    
    st.success(get_lang_text('preset_scores_loaded').format(grade=grade))
    st.rerun()

# 加载分值权重配置
score_weights_config = config.score_weights

# 计算合规分数
@handle_errors(default_return_value=0, error_type=BusinessLogicError)
def calculate_compliance_score(responses, question_type, sub_responses=None, score_weight=1):
    """计算合规分数
    
    根据问题类型和回答情况计算加权合规分数，支持三种问题类型：XO（是/否）、PJ（等级评分）和PW（多选）。
    
    参数:
        responses (int): 问题的主回答得分，范围0-4（XO类型：4表示"是"，0表示"否"）
        question_type (str): 问题类型，必须是以下之一："XO"（是/否）、"PJ"（等级评分）、"PW"（多选）
        sub_responses (dict, optional): PW类型问题的子问题回答情况，键为子问题ID，值为布尔值表示是否选中
        score_weight (float, optional): 问题的权重系数，默认为1
        
    返回:
        float: 计算得到的加权合规分数
        
    计算逻辑:
        - XO类型: 根据是否为"是"(4分)获取对应基础分数
        - PJ类型: 直接从配置中获取对应评分的基础分数
        - PW类型: 计算子问题选中比例，转换为基础分数
        - 所有类型最终都会乘以权重系数得到最终分数
        
    示例:
        # XO类型问题，回答"是"，权重2
        calculate_compliance_score(4, "XO", None, 2)  # 返回 2 * 100 = 200
        
        # PW类型问题，5个子问题选中3个，权重10
        calculate_compliance_score(0, "PW", {"sub_1": True, "sub_2": True, "sub_3": True, "sub_4": False, "sub_5": False}, 10)  # 返回 10 * (3/5) = 6
    """
    # 边界检查：score_weight必须为正数
    if score_weight <= 0:
        return 0
    
    # 边界检查：question_type必须为有效类型
    if question_type not in ["XO", "PJ", "PW"]:
        return 0
    
    if question_type == "PW":
        # 边界检查：sub_responses必须为非空字典
        if not isinstance(sub_responses, dict) or not sub_responses:
            return 0
        # PW类型：计算子问题选中比例（转换为百分比）
        sub_scores = [100 if bool(v) else 0 for v in sub_responses.values()]
        base_score = sum(sub_scores) / len(sub_scores) if sub_scores else 0
    elif question_type == "XO":
        # 边界检查：responses必须为0-4的整数
        if not isinstance(responses, (int, float)) or responses < 0 or responses > 4:
            return 0
        type_base = score_weights_config['question_type_base_scores']
        # XO类型："是"(4分)对应yes的基础分，否则对应no的基础分
        base_score = type_base['XO'].get("yes" if responses == 4 else "no", 0)
    elif question_type == "PJ":
        # 边界检查：responses必须为0-4的整数
        if not isinstance(responses, (int, float)) or responses < 0 or responses > 4:
            return 0
        type_base = score_weights_config['question_type_base_scores']
        # PJ类型：直接获取对应评分的基础分
        base_score = type_base['PJ'].get(int(responses), 0)
    else:
        # 无效的问题类型，返回0
        return 0
    
    # 边界检查：base_score必须在0-100之间
    if base_score < 0 or base_score > 100:
        return 0
    
    # 计算加权分数（基础分数 / 100 * 权重）
    weighted_score = (base_score / 100) * score_weight
    
    # 确保返回值为正数
    return max(0, weighted_score)

# 计算总分（1000分制）
@handle_errors(default_return_value=0, error_type=BusinessLogicError)
def calculate_total_score(section_scores):
    """计算1000分制总分"""
    if not section_scores:
        return 0
    # 直接返回所有章节得分之和
    return sum(section_scores.values())

# 生成雷达图
@handle_errors(default_return_value=None, error_type=BusinessLogicError)
def create_radar_chart(section_scores, language=None):
    """生成雷达图
    
    根据各章节得分生成雷达图，显示ISO 55001评估各维度的合规情况。
    
    参数:
        section_scores (dict): 各章节得分，键为章节名称，值为得分（0-1000分制）
        language (str, optional): 显示语言，'en'或'zh'，默认使用st.session_state.language或'zh'
        
    返回:
        plotly.graph_objects.Figure: 生成的雷达图对象，可用于Streamlit显示或导出
        None: 如果section_scores为空或无效
        
    图表特点:
        - 雷达图显示各章节的得分百分比（相对于该章节满分）
        - 百分比精确到1位小数，99.99%会显示为100.0% 
        - 支持中英文显示
        - 提供清晰的颜色区分和标注
        
    示例:
        section_scores = {"组织环境": 85, "领导作用": 75, "策划": 90}
        fig = create_radar_chart(section_scores)
        if fig:
            st.plotly_chart(fig)
    """
    if not section_scores:
        return None
    section_weights = score_weights_config['section_weights']
    # 获取章节名称和分数
    categories = []  # 雷达图显示的章节名称列表
    values = []  # 雷达图显示的得分百分比列表
    
    # 获取当前语言环境
    current_language = language or get_language()
    
    for section, score in section_scores.items():
        # 根据当前语言环境获取章节显示名称
        section_name = get_section_title(section, current_language)
        categories.append(section_name)
        max_score = section_weights.get(section, 100)  # 章节满分（默认为100分）
        # 计算雷达图显示的百分比值
        # 1. 当得分与满分非常接近（差值小于1e-6）时，直接显示为100（处理浮点数精度问题）
        # 2. 否则计算百分比并四舍五入到1位小数（如99.99%显示为100.0%，提高可读性）
        # 3. 确保max_score不为0，避免除零错误
        value = 100 if abs(score - max_score) < 1e-6 else (round(score / max_score * 100, 1) if max_score else 0)
        values.append(value)
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=categories,
        fill='toself',
        name=f"{get_lang_text('compliance_score')} (%)",
        line_color='#4CAF50'
    ))
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                tickfont=dict(size=12),
                gridcolor='#f0f2f6',
                tickformat='.1f%'
            ),
            angularaxis=dict(
                tickfont=dict(size=12),
                gridcolor='#f0f2f6'
            ),
            bgcolor='white'
        ),
        showlegend=True,
        paper_bgcolor='white',
        plot_bgcolor='white',
        margin=dict(t=30, b=30, l=30, r=30)
    )
    return fig

# 获取PDF报告所需的字体
@handle_errors(default_return_value=("STSong-Light", "STSong-Light"), error_type=ExportError)
def get_pdf_fonts():
    """获取PDF报告所需的字体，实现多级降级策略
    
    字体处理逻辑:
        1. 优先使用指定的中文字体文件（simsun.ttc和simhei.ttf）
        2. 如果字体文件不存在，尝试使用系统兼容字体
        3. 最终降级至ReportLab内置中文字体
        
    返回:
        tuple: (main_font, bold_font) - 主字体和粗体字体名称
    """
    # 字体配置和备选列表
    font_dir = Path(__file__).parent / "fonts"
    chinese_fonts = {
        'main': ('SimSun', font_dir / "simsun.ttc"),
        'bold': ('SimHei', font_dir / "simhei.ttf")
    }
    fallback_fonts = ['STSong-Light', 'WenQuanYi Micro Hei', 'Heiti TC']
    
    # 英文环境处理：直接返回配置的默认字体
    if st.session_state.language != 'zh':
        return config.general['font_en'], config.general['font_en_bold']
    
    # 中文环境处理：尝试多级降级策略
    # 1. 优先使用项目提供的中文字体文件
    simsun_exists = chinese_fonts['main'][1].exists()
    simhei_exists = chinese_fonts['bold'][1].exists()
    
    if simsun_exists and simhei_exists:
        pdfmetrics.registerFont(TTFont(chinese_fonts['main'][0], str(chinese_fonts['main'][1])))
        pdfmetrics.registerFont(TTFont(chinese_fonts['bold'][0], str(chinese_fonts['bold'][1])))
        return config.general['font_zh'], config.general['font_zh_bold']
    
    # 2. 尝试使用系统兼容字体
    logging.warning(f"中文环境下缺少字体文件: simsun.ttc={'存在' if simsun_exists else '不存在'}, simhei.ttf={'存在' if simhei_exists else '不存在'}")
    
    for font_name in fallback_fonts:
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_name))
            return font_name, font_name
        except Exception:
            continue
    
    # 3. 最终降级至ReportLab内置中文字体
    return 'STSong-Light', 'STSong-Light'

@handle_errors(default_return_value=None, error_type=ExportError)
def create_pdf_report(section_scores, questionnaire, responses, sub_responses):
    """生成PDF报告
    
    生成包含ISO 55001评估结果的PDF报告，包括：
    - 整体评分概览
    - 各章节详细得分
    - 雷达图可视化
    - 各问题的具体回答和得分情况
    
    参数:
        section_scores (dict): 各章节得分，键为章节名称，值为得分（0-1000分制）
        questionnaire (dict): 评估问题字典，包含各章节的问题详情
        responses (dict): 问题的主回答，键为问题ID，值为得分
        sub_responses (dict): PW类型问题的子回答，键为子问题ID，值为布尔值表示是否选中
        
    返回:
        io.BytesIO: 生成的PDF文件字节流，可直接用于下载或保存
        None: 如果缺少必要数据或生成过程中出现错误
        
    报告特点:
        - 支持中英文双语
        - 自动处理中文字体，确保中文显示正常
        - 包含完整的评分结果和可视化图表
        - 专业的排版和样式设计
        
    字体处理逻辑:
        1. 优先使用指定的中文字体文件（simsun.ttc和simhei.ttf）
        2. 如果字体文件不存在，尝试使用系统兼容字体
        3. 最终降级至ReportLab内置中文字体
        
    示例:
        pdf_buffer = create_pdf_report(section_scores, questionnaire, responses, sub_responses)
        if pdf_buffer:
            with open("report.pdf", "wb") as f:
                f.write(pdf_buffer.getvalue())
    """
    if not section_scores or not questionnaire:
        logging.error("生成PDF报告失败：缺少必要数据")
        return None
    
    # 获取字体
    main_font, bold_font = get_pdf_fonts()
    
    # 创建PDF文档
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          leftMargin=50,
                          rightMargin=50,
                          topMargin=50,
                          bottomMargin=50)
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=bold_font,
        fontSize=28,
        spaceAfter=30,
        alignment=1,
        textColor=colors.HexColor('#2E4053')
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontName=bold_font,
        fontSize=20,
        spaceAfter=15,
        textColor=colors.HexColor('#2874A6')
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontName=bold_font,
        fontSize=16,
        spaceAfter=12,
        textColor=colors.HexColor('#3498DB')
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=main_font,
        fontSize=12,
        spaceAfter=8,
        leading=16,
        textColor=colors.black
    )
    
    elements = []
    
    # 添加标题
    title = config.lang_zh['app_title'] if st.session_state.language == 'zh' else config.lang_en['app_title']
    # 组合第一页内容
    first_page_content = []
    first_page_content.append(Paragraph(title, title_style))
    # 总体评分
    total_score = sum(section_scores.values()) if section_scores else 0
    score_style = ParagraphStyle(
        'ScoreStyle',
        parent=heading2_style,
        fontSize=24,
        textColor=colors.HexColor('#27AE60')
    )
    overall_score = config.lang_zh['overall_score'] if st.session_state.language == 'zh' else config.lang_en['overall_score']
    first_page_content.append(Paragraph(f"{overall_score}{total_score:.1f}", score_style))
    first_page_content.append(Spacer(1, 10))
    # 雷达图
    radar_chart = create_radar_chart(section_scores)
    if radar_chart:
        try:
            img_data = radar_chart.to_image(format="png")
            img = Image(io.BytesIO(img_data), width=6.8*inch, height=5.0*inch)  # 宽高
            first_page_content.append(img)
            first_page_content.append(Spacer(1, 10))
        except Exception as e:
            logging.error(f"添加雷达图到PDF失败: {str(e)}")
    # 各要素得分详情表格
    first_page_content.append(Paragraph(
        config.lang_zh['element_scores_detail'] if st.session_state.language == 'zh' else config.lang_en['element_scores_detail'],
        heading2_style
    ))
    first_page_content.append(Spacer(1, 5))
    headers = [config.lang_zh['element'] if st.session_state.language == 'zh' else config.lang_en['element'], config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']]
    data = [headers]
    for section, score in section_scores.items():
        section_id = questionnaire[section]['name']['zh'] if st.session_state.language == 'zh' else questionnaire[section]['name']['en']
        data.append([section_id, f"{score:.1f}"])
    col_widths = [doc.width/2.0, doc.width/2.0]
    table = Table(data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E4053')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), bold_font),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9F9')),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), main_font),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#D5D8DC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9F9')]),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    first_page_content.append(table)
    first_page_content.append(Spacer(1, 10))
    elements.append(KeepTogether(first_page_content))
    elements.append(PageBreak())
    
    # 添加详细评估结果
    elements.append(Paragraph(
        config.lang_zh['detailed_assessment_results'] if st.session_state.language == 'zh' else config.lang_en['detailed_assessment_results'],
        heading2_style
    ))
    elements.append(Spacer(1, 15))
    
    score_labels = {
        'zh': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        },
        'en': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        }
    }
    
    for section, section_data in questionnaire.items():
        section_id = section_data['name']['zh'] if st.session_state.language == 'zh' else section_data.get('id', section)
        elements.append(Paragraph(section_id, heading3_style))
        
        for q_id, question in section_data.get('questions', {}).items():
            key = f"{section}_{q_id}"
            score = responses.get(key, 0)
            question_weight = score_weights_config['question_weights'][section].get(q_id, 1)
            # 计算实际得分
            if question['type'] == "PW":
                # 准备PW题的子响应数据
                sub_questions = question.get("sub_questions", {}).get(st.session_state.language, [])
                pw_sub_responses = {}
                for i in range(1, len(sub_questions) + 1):
                    sub_key = f"{key}_sub_{i}"
                    pw_sub_responses[sub_key] = st.session_state.sub_responses.get(sub_key, False)
                actual_score = calculate_compliance_score(score, question['type'], pw_sub_responses, question_weight)
            else:
                actual_score = calculate_compliance_score(score, question['type'], None, question_weight)
            # 创建问题样式
            question_style = ParagraphStyle(
                'QuestionStyle',
                parent=normal_style,
                fontSize=13,
                textColor=colors.HexColor('#2C3E50'),
                spaceAfter=5
            )
            # 创建得分样式
            score_style = ParagraphStyle(
                'ScoreStyle',
                parent=normal_style,
                fontSize=12,
                textColor=colors.HexColor('#E74C3C'),
                spaceAfter=8
            )
            # 添加问题描述和得分
            question_text = config.lang_zh['question'] if st.session_state.language == 'zh' else config.lang_en['question']
            type_text = config.lang_zh['type'] if st.session_state.language == 'zh' else config.lang_en['type']
            score_text = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
            weight_text = config.lang_zh['weight'] if st.session_state.language == 'zh' else config.lang_en['weight']
            description = get_translated_text(question["description"], st.session_state.language)
            elements.append(Paragraph(f"{question_text}{description}", question_style))
            elements.append(Paragraph(f"{type_text}{question['type']}", normal_style))
            elements.append(Paragraph(f"{score_text}{actual_score:.1f}", score_style))
            elements.append(Paragraph(f"{weight_text}{question_weight}", normal_style))
            
            # 如果是多选题，添加子问题得分
            if question['type'] == "PW" and "sub_questions" in question:
                sub_questions = question["sub_questions"].get(st.session_state.language, [])
                sub_scores_text = config.lang_zh['sub_question_scores'] if st.session_state.language == 'zh' else config.lang_en['sub_question_scores']
                elements.append(Paragraph(sub_scores_text, normal_style))
                sub_count = len(sub_questions)
                question_weight = score_weights_config['question_weights'][section].get(q_id, 1)
                for i, sub_q in enumerate(sub_questions, 1):
                    sub_key = f"{key}_sub_{i}"
                    sub_score = st.session_state.sub_responses.get(sub_key, False)
                    yes_text = config.lang_zh['answer_yes'] if st.session_state.language == 'zh' else config.lang_en['answer_yes']
                    no_text = config.lang_zh['answer_no'] if st.session_state.language == 'zh' else config.lang_en['answer_no']
                    # 子问题得分为题目权重/子项数
                    sub_score_value = question_weight / sub_count if sub_score else 0
                    if st.session_state.language == 'en':
                        elements.append(Paragraph(f"- {sub_q}: {yes_text if sub_score else no_text} ({sub_score_value:.1f})", normal_style))
                    else:
                        elements.append(Paragraph(f"- {sub_q}: {yes_text if sub_score else no_text}（{sub_score_value:.1f}）", normal_style))
            
            elements.append(Spacer(1, 15))
    
    # 生成PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

# 创建Markdown报告
@handle_errors(default_return_value=None, error_type=ExportError)
def create_markdown_report(section_scores, questionnaire, responses, sub_responses):
    """生成Markdown格式报告
    
    生成包含ISO 55001评估结果的Markdown报告，包括：
    - 整体评分概览
    - 各章节详细得分（表格形式）
    - 各问题的具体回答和得分情况
    
    参数:
        section_scores (dict): 各章节得分，键为章节名称，值为得分（0-1000分制）
        questionnaire (dict): 评估问题字典，包含各章节的问题详情
        responses (dict): 问题的主回答，键为问题ID，值为得分
        sub_responses (dict): PW类型问题的子回答，键为子问题ID，值为布尔值表示是否选中
        
    返回:
        str: 生成的Markdown报告内容
        None: 如果缺少必要数据或生成过程中出现错误
    """
    if not section_scores or not questionnaire:
        logging.error("生成Markdown报告失败：缺少必要数据")
        return None
    
    # 创建Markdown内容
    md_content = []
    
    # 报告标题
    title = config.lang_zh['app_title'] if st.session_state.language == 'zh' else config.lang_en['app_title']
    md_content.append(f"# {title}")
    md_content.append("")
    
    # 报告日期
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_text = config.lang_zh['report_generate_time'] if st.session_state.language == 'zh' else config.lang_en['report_generate_time']
    md_content.append(f"{date_text}: {report_date}")
    md_content.append("")
    
    # 总体评分
    total_score = sum(section_scores.values()) if section_scores else 0
    overall_score_text = config.lang_zh['overall_score'] if st.session_state.language == 'zh' else config.lang_en['overall_score']
    md_content.append(f"## {overall_score_text}")
    md_content.append(f"### {total_score:.1f}")
    md_content.append("")
    
    # 各章节得分详情表格
    element_scores_text = config.lang_zh['element_scores_detail'] if st.session_state.language == 'zh' else config.lang_en['element_scores_detail']
    md_content.append(f"## {element_scores_text}")
    md_content.append("")
    
    # 表格表头
    element_text = config.lang_zh['element'] if st.session_state.language == 'zh' else config.lang_en['element']
    score_text = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
    md_content.append(f"| {element_text} | {score_text} |")
    md_content.append("|-------------|--------|")
    
    # 表格数据
    for section, score in section_scores.items():
        section_id = questionnaire[section]['name']['zh'] if st.session_state.language == 'zh' else questionnaire[section]['name']['en']
        md_content.append(f"| {section_id} | {score:.1f} |")
    md_content.append("")
    
    # 详细评估结果
    detailed_results_text = config.lang_zh['detailed_assessment_results'] if st.session_state.language == 'zh' else config.lang_en['detailed_assessment_results']
    md_content.append(f"## {detailed_results_text}")
    md_content.append("")
    
    # 得分标签
    score_labels = {
        'zh': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        },
        'en': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        }
    }
    
    # 遍历各章节和问题
    for section, section_data in questionnaire.items():
        section_name = section_data['name']['zh'] if st.session_state.language == 'zh' else section_data['name']['en']
        md_content.append(f"### {section_name}")
        md_content.append("")
        
        for q_id, question in section_data.get('questions', {}).items():
            key = f"{section}_{q_id}"
            score = responses.get(key, 0)
            question_weight = score_weights_config['question_weights'][section].get(q_id, 1)
            
            # 计算实际得分
            if question['type'] == "PW":
                # 准备PW题的子响应数据
                sub_questions = question.get("sub_questions", {}).get(st.session_state.language, [])
                pw_sub_responses = {}
                for i in range(1, len(sub_questions) + 1):
                    sub_key = f"{key}_sub_{i}"
                    pw_sub_responses[sub_key] = sub_responses.get(sub_key, False)
                actual_score = calculate_compliance_score(score, question['type'], pw_sub_responses, question_weight)
            else:
                actual_score = calculate_compliance_score(score, question['type'], None, question_weight)
            
            # 添加问题描述
            description = get_translated_text(question["description"], st.session_state.language)
            question_text = config.lang_zh['question'] if st.session_state.language == 'zh' else config.lang_en['question']
            md_content.append(f"#### {description}")
            
            # 添加问题信息
            type_text = config.lang_zh['type'] if st.session_state.language == 'zh' else config.lang_en['type']
            score_text = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
            weight_text = config.lang_zh['weight'] if st.session_state.language == 'zh' else config.lang_en['weight']
            assessment_text = config.lang_zh['assessment'] if st.session_state.language == 'zh' else config.lang_en['assessment']
            
            # 根据问题类型获取评估结果
            if question["type"] != "XO":
                assessment = score_labels[st.session_state.language][round(score)]
            else:
                assessment = get_lang_text('answer_yes') if score == 4 else get_lang_text('answer_no')
            
            md_content.append(f"- {type_text}: {question['type']}")
            md_content.append(f"- {score_text}: {actual_score:.1f}")
            md_content.append(f"- {weight_text}: {question_weight}")
            md_content.append(f"- {assessment_text}: {assessment}")
            
            # 如果是多选题，添加子问题得分
            if question['type'] == "PW" and "sub_questions" in question:
                sub_questions = question["sub_questions"].get(st.session_state.language, [])
                sub_scores_text = config.lang_zh['sub_question_scores'] if st.session_state.language == 'zh' else config.lang_en['sub_question_scores']
                md_content.append(f"- {sub_scores_text}:")
                
                for i, sub_q in enumerate(sub_questions, 1):
                    sub_key = f"{key}_sub_{i}"
                    sub_score = sub_responses.get(sub_key, False)
                    status = "√" if sub_score else "×"
                    md_content.append(f"  - {status} {sub_q}")
            
            md_content.append("")
    
    # 合并所有内容
    return "\n".join(md_content)

# 创建Excel报告
@handle_errors(default_return_value=None, error_type=ExportError)
def create_excel_report(section_scores, questionnaire, responses, sub_responses):
    """生成Excel格式报告
    
    生成包含ISO 55001评估结果的Excel报告，包括：
    - 整体评分概览
    - 各章节详细得分（表格形式）
    - 各问题的具体回答和得分情况
    
    参数:
        section_scores (dict): 各章节得分，键为章节名称，值为得分（0-1000分制）
        questionnaire (dict): 评估问题字典，包含各章节的问题详情
        responses (dict): 问题的主回答，键为问题ID，值为得分
        sub_responses (dict): PW类型问题的子回答，键为子问题ID，值为布尔值表示是否选中
        
    返回:
        io.BytesIO: 生成的Excel报告的字节流
        None: 如果缺少必要数据或生成过程中出现错误
    """
    if not section_scores or not questionnaire:
        logging.error("生成Excel报告失败：缺少必要数据")
        return None
    
    try:
        # 导入openpyxl库
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from openpyxl.cell.cell import MergedCell
        from openpyxl.utils.dataframe import dataframe_to_rows
        import pandas as pd
        import io
        
        # 创建工作簿
        wb = Workbook()
        
        # 创建报告数据工作表
        ws = wb.active
        ws.title = "Assessment Results"
        
        # 设置页面边距
        ws.page_margins.top = 1
        ws.page_margins.bottom = 1
        ws.page_margins.left = 1
        ws.page_margins.right = 1
        
        # 报告标题
        ws.merge_cells('A1:G1')
        title_cell = ws['A1']
        title = config.lang_zh['app_title'] if st.session_state.language == 'zh' else config.lang_en['app_title']
        title_cell.value = title
        title_cell.font = Font(bold=True, size=24)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 报告日期
        ws.merge_cells('A2:G2')
        date_cell = ws['A2']
        report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        date_text = config.lang_zh['report_generate_time'] if st.session_state.language == 'zh' else config.lang_en['report_generate_time']
        date_cell.value = f"{date_text}: {report_date}"
        date_cell.font = Font(size=12)
        date_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 总体评分
        ws.merge_cells('A3:G3')
        total_score_cell = ws['A3']
        total_score = sum(section_scores.values()) if section_scores else 0
        overall_score_text = config.lang_zh['overall_score'] if st.session_state.language == 'zh' else config.lang_en['overall_score']
        total_score_cell.value = f"{overall_score_text}: {total_score:.1f}/1000"
        total_score_cell.font = Font(bold=True, size=16)
        total_score_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 添加空行
        ws.append([])
        ws.append([])
        
        # 各章节得分表格
        section_scores_text = config.lang_zh['element_scores_detail'] if st.session_state.language == 'zh' else config.lang_en['element_scores_detail']
        ws['A6'] = section_scores_text
        ws['A6'].font = Font(bold=True, size=14)
        
        # 章节得分表格标题
        ws['A7'] = config.lang_zh['element'] if st.session_state.language == 'zh' else config.lang_en['element']
        ws['B7'] = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
        
        # 设置表头样式
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color='D9D9D9', end_color='D9D9D9', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        for cell in ws[7]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # 填充章节得分数据
        row = 8
        for section, score in section_scores.items():
            section_name = questionnaire[section]['name'][st.session_state.language]
            ws[f'A{row}'] = section_name
            ws[f'B{row}'] = score
            
            # 设置数据样式
            ws[f'A{row}'].border = thin_border
            ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='center')
            ws[f'B{row}'].border = thin_border
            ws[f'B{row}'].alignment = Alignment(horizontal='center', vertical='center')
            ws[f'B{row}'].number_format = '0.0'
            
            row += 1
        
        # 自动调整列宽
        for column in ws.columns:
            max_length = 0
            column_letter = None
            
            # 找到非合并单元格来获取列字母
            for cell in column:
                if not isinstance(cell, MergedCell):
                    column_letter = cell.column_letter
                    break
            
            if column_letter:
                for cell in column:
                    try:
                        if not isinstance(cell, MergedCell) and cell.value is not None:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        # 创建问题详情工作表
        ws_details = wb.create_sheet(title="Detailed Results")
        
        # 问题详情表格标题
        ws_details['A1'] = config.lang_zh['detailed_assessment_results'] if st.session_state.language == 'zh' else config.lang_en['detailed_assessment_results']
        ws_details['A1'].font = Font(bold=True, size=14)
        
        # 问题详情表头
        ws_details['A3'] = config.lang_zh['element'] if st.session_state.language == 'zh' else config.lang_en['element']
        ws_details['B3'] = config.lang_zh['question'] if st.session_state.language == 'zh' else config.lang_en['question']
        ws_details['C3'] = config.lang_zh['type'] if st.session_state.language == 'zh' else config.lang_en['type']
        ws_details['D3'] = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
        ws_details['E3'] = config.lang_zh['weight'] if st.session_state.language == 'zh' else config.lang_en['weight']
        
        # 设置表头样式
        for cell in ws_details[3]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # 填充问题详情数据
        row = 4
        for section, section_data in questionnaire.items():
            section_name = section_data['name'][st.session_state.language]
            for q_id, question in section_data.get('questions', {}).items():
                key = f"{section}_{q_id}"
                score = responses.get(key, 0)
                question_weight = score_weights_config['question_weights'][section].get(q_id, 1)
                
                ws_details[f'A{row}'] = section_name
                ws_details[f'B{row}'] = get_translated_text(question["description"], st.session_state.language)
                ws_details[f'C{row}'] = question["type"]
                ws_details[f'D{row}'] = score
                ws_details[f'E{row}'] = question_weight
                
                # 设置数据样式
                for col in ['A', 'B', 'C', 'D', 'E']:
                    ws_details[f'{col}{row}'].border = thin_border
                    ws_details[f'{col}{row}'].alignment = Alignment(horizontal='left', vertical='center')
                    
                if col == 'D' or col == 'E':
                    ws_details[f'{col}{row}'].alignment = Alignment(horizontal='center', vertical='center')
                
                row += 1
        
        # 自动调整详情页列宽
        for column in ws_details.columns:
            max_length = 0
            column_letter = None
            
            # 找到非合并单元格来获取列字母
            for cell in column:
                if not isinstance(cell, MergedCell):
                    column_letter = cell.column_letter
                    break
            
            if column_letter:
                for cell in column:
                    try:
                        if not isinstance(cell, MergedCell) and cell.value is not None:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws_details.column_dimensions[column_letter].width = adjusted_width
        
        # 保存到字节流
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        logging.error(f"生成Excel报告失败：{str(e)}")
        logging.error(traceback.format_exc())
        return None

# 创建DOCX报告
@handle_errors(default_return_value=None, error_type=ExportError)
def create_doc_report(section_scores, questionnaire, responses, sub_responses):
    """生成DOCX格式报告
    
    生成包含ISO 55001评估结果的DOCX报告，包括：
    - 整体评分概览
    - 各章节详细得分（表格形式）
    - 各问题的具体回答和得分情况
    
    参数:
        section_scores (dict): 各章节得分，键为章节名称，值为得分（0-1000分制）
        questionnaire (dict): 评估问题字典，包含各章节的问题详情
        responses (dict): 问题的主回答，键为问题ID，值为得分
        sub_responses (dict): PW类型问题的子回答，键为子问题ID，值为布尔值表示是否选中
        
    返回:
        io.BytesIO: 生成的DOCX报告的字节流
        None: 如果缺少必要数据或生成过程中出现错误
    """
    if not section_scores or not questionnaire:
        logging.error("生成DOCX报告失败：缺少必要数据")
        return None
    
    # 创建新文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 报告标题
    title = config.lang_zh['app_title'] if st.session_state.language == 'zh' else config.lang_en['app_title']
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 报告日期
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_text = config.lang_zh['report_generate_time'] if st.session_state.language == 'zh' else config.lang_en['report_generate_time']
    date_para = doc.add_paragraph(f"{date_text}: {report_date}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 总体评分
    total_score = sum(section_scores.values()) if section_scores else 0
    overall_score_text = config.lang_zh['overall_score'] if st.session_state.language == 'zh' else config.lang_en['overall_score']
    doc.add_heading(overall_score_text, level=1)
    total_score_para = doc.add_paragraph(f"{total_score:.1f}")
    total_score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_score_run = total_score_para.runs[0]
    total_score_run.font.size = Pt(36)
    total_score_run.font.bold = True
    doc.add_paragraph()
    
    # 各章节得分详情表格
    element_scores_text = config.lang_zh['element_scores_detail'] if st.session_state.language == 'zh' else config.lang_en['element_scores_detail']
    doc.add_heading(element_scores_text, level=1)
    
    # 表格表头
    element_text = config.lang_zh['element'] if st.session_state.language == 'zh' else config.lang_en['element']
    score_text = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
    
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.autofit = False
    table.columns[0].width = Inches(4)
    table.columns[1].width = Inches(2)
    
    # 设置表头样式
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = element_text
    hdr_cells[1].text = score_text
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 表格数据
    for section, score in section_scores.items():
        section_name = questionnaire[section]['name']['zh'] if st.session_state.language == 'zh' else questionnaire[section]['name']['en']
        row_cells = table.add_row().cells
        row_cells[0].text = section_name
        row_cells[1].text = f"{score:.1f}"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 得分标签
    score_labels = {
        'zh': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        },
        'en': {
            0: get_lang_text('not_implemented'),
            1: get_lang_text('preliminary_implementation'),
            2: get_lang_text('partial_implementation'),
            3: get_lang_text('mostly_implemented'),
            4: get_lang_text('fully_implemented')
        }
    }
    
    # 详细评估结果
    detailed_results_text = config.lang_zh['detailed_assessment_results'] if st.session_state.language == 'zh' else config.lang_en['detailed_assessment_results']
    doc.add_heading(detailed_results_text, level=1)
    
    # 遍历各章节和问题
    for section, section_data in questionnaire.items():
        section_name = section_data['name']['zh'] if st.session_state.language == 'zh' else section_data['name']['en']
        doc.add_heading(section_name, level=2)
        
        for q_id, question in section_data.get('questions', {}).items():
            key = f"{section}_{q_id}"
            score = responses.get(key, 0)
            question_weight = score_weights_config['question_weights'][section].get(q_id, 1)
            
            # 计算实际得分
            if question['type'] == "PW":
                # 准备PW题的子响应数据
                sub_questions = question.get("sub_questions", {}).get(st.session_state.language, [])
                pw_sub_responses = {}
                for i in range(1, len(sub_questions) + 1):
                    sub_key = f"{key}_sub_{i}"
                    pw_sub_responses[sub_key] = sub_responses.get(sub_key, False)
                actual_score = calculate_compliance_score(score, question['type'], pw_sub_responses, question_weight)
            else:
                actual_score = calculate_compliance_score(score, question['type'], None, question_weight)
            
            # 添加问题描述
            description = get_translated_text(question["description"], st.session_state.language)
            doc.add_heading(description, level=3)
            
            # 根据问题类型获取评估结果
            if question["type"] != "XO":
                assessment = score_labels[st.session_state.language][round(score)]
            else:
                assessment = get_lang_text('answer_yes') if score == 4 else get_lang_text('answer_no')
            
            type_text = config.lang_zh['type'] if st.session_state.language == 'zh' else config.lang_en['type']
            score_text = config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score']
            weight_text = config.lang_zh['weight'] if st.session_state.language == 'zh' else config.lang_en['weight']
            assessment_text = config.lang_zh['assessment'] if st.session_state.language == 'zh' else config.lang_en['assessment']
            
            # 添加问题信息
            doc.add_paragraph(f"{type_text}: {question['type']}", style='List Bullet')
            doc.add_paragraph(f"{score_text}: {actual_score:.1f}", style='List Bullet')
            doc.add_paragraph(f"{weight_text}: {question_weight}", style='List Bullet')
            doc.add_paragraph(f"{assessment_text}: {assessment}", style='List Bullet')
            
            # 如果是多选题，添加子问题得分
            if question['type'] == "PW" and "sub_questions" in question:
                sub_questions = question["sub_questions"].get(st.session_state.language, [])
                sub_scores_text = config.lang_zh['sub_question_scores'] if st.session_state.language == 'zh' else config.lang_en['sub_question_scores']
                doc.add_paragraph(f"{sub_scores_text}:", style='List Bullet')
                
                for i, sub_q in enumerate(sub_questions, 1):
                    sub_key = f"{key}_sub_{i}"
                    sub_score = sub_responses.get(sub_key, False)
                    status = "√" if sub_score else "×"
                    doc.add_paragraph(f"{status} {sub_q}", style='List Bullet 2')
            
            doc.add_paragraph()
    
    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 初始化数据库
init_db()

# 设置页面配置
st.set_page_config(
    page_title="ISO 55001 评估工具",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded"
)

# 加载外部CSS文件
with open('style.css', encoding='utf-8') as f:
    st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

@handle_errors(default_return_value='', error_type=ConfigError)
def get_translated_text(text_dict, lang='zh'):
    """获取翻译文本"""
    if isinstance(text_dict, str):
        return text_dict
    if isinstance(text_dict, dict):
        return text_dict.get(lang, text_dict.get('zh', ''))
    return ''

@handle_errors(default_return_value='**未知章节**', error_type=ConfigError)
def get_display_section_title(section_data, lang='zh'):
    """获取章节标题"""
    if section_data is None or not isinstance(section_data, dict):
        section_data = {}
    name = get_translated_text(section_data.get('name', section_data.get('id', '')), lang)
    if not name:
        name = '未知章节'
    if lang == 'zh':
        return f"**{name}**"
    else:
        return f"**{name.title()}**"

@handle_errors(default_return_value=None, error_type=BusinessLogicError)
def main():
    """主函数"""
    try:
        # 初始化会话状态
        init_session_state()
        
        # 加载评估问题
        try:
            questionnaire = load_questionnaire()
        except Exception as e:
            st.error(f"{get_lang_text('error_loading_questions')}: {str(e)}")
            return

        # 添加侧边栏
        with st.sidebar:
            # 简化语言切换为两个按钮
            col1, col2 = st.columns(2)
            with col1:
                if st.button(config.lang_zh['zh_button'] if st.session_state.language == 'zh' else config.lang_en['zh_button'], type="primary" if st.session_state.language == 'zh' else "secondary"):
                    st.session_state.language = 'zh'
                    st.rerun()
            with col2:
                if st.button(config.lang_zh['en_button'] if st.session_state.language == 'en' else config.lang_en['en_button'], type="primary" if st.session_state.language == 'en' else "secondary"):
                    st.session_state.language = 'en'
                    st.rerun()
            
            # 添加标题
            st.title(config.lang_zh['app_title'] if st.session_state.language == 'zh' else config.lang_en['app_title'])
            
            st.markdown("---")
            
            # 添加保存和加载按钮
            col1, col2 = st.columns(2)
            with col1:
                save_text = config.lang_zh['save_progress'] if st.session_state.language == 'zh' else config.lang_en['save_progress']
                if st.button(save_text, key="save_button"):
                    try:
                        save_assessment_results(st.session_state.responses, st.session_state.sub_responses)
                        st.session_state.last_save_time = datetime.now()
                        st.success(config.lang_zh['progress_saved'] if st.session_state.language == 'zh' else config.lang_en['progress_saved'])
                    except Exception as e:
                        st.error(f"{config.lang_zh['error_saving_progress']} {str(e)}" if st.session_state.language == 'zh' else f"{config.lang_en['error_saving_progress']} {str(e)}")
            
            with col2:
                load_text = config.lang_zh['load_progress'] if st.session_state.language == 'zh' else config.lang_en['load_progress']
                if st.button(load_text, key="load_button"):
                    try:
                        responses, sub_responses = load_latest_assessment_results()
                        st.session_state.responses = responses
                        st.session_state.sub_responses = sub_responses
                        st.session_state.force_refresh = True
                        st.success(config.lang_zh['last_progress_loaded'] if st.session_state.language == 'zh' else config.lang_en['last_progress_loaded'])
                        st.rerun()
                    except Exception as e:
                        st.error(f"{config.lang_zh['error_loading_progress']} {str(e)}" if st.session_state.language == 'zh' else f"{config.lang_en['error_loading_progress']} {str(e)}")
            
            # 显示上次保存时间
            last_save_text = config.lang_zh['last_saved'] if st.session_state.language == 'zh' else config.lang_en['last_saved']
            st.markdown(f"{last_save_text}{st.session_state.last_save_time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            st.markdown("---")
            if st.session_state.language == 'en':
                st.markdown(f"""
                #### {get_lang_text('question_types')}
                - PJ：{get_lang_text('question_type_pj')}
                - XO：{get_lang_text('question_type_xo')}
                - PW：{get_lang_text('question_type_pw')}
                """)
            else:
                st.markdown(f"""
                #### {get_lang_text('question_types')}
                - PJ：{get_lang_text('question_type_pj')}
                - XO：{get_lang_text('question_type_xo')}
                - PW：{get_lang_text('question_type_pw')}
                """)
            
            # 添加预设分值按钮
            st.markdown("---")
            if st.session_state.language == 'en':
                st.markdown(f"#### {get_lang_text('preset_scores_title')}")
                preset_score_text = get_lang_text('load_preset_scores')
            else:
                st.markdown(f"#### {get_lang_text('preset_scores')}")
                preset_score_text = get_lang_text('load_preset_scores')
            
            st.markdown(preset_score_text)
            cols = st.columns(5)
            grades = ['A', 'B', 'C', 'D', 'E']
            for i, grade in enumerate(grades):
                with cols[i]:
                    if st.button(grade, key=f"preset_{grade}"):
                        apply_preset_scores(grade)

        # 创建选项卡
        tab_titles = [get_lang_text('system_assessment'), get_lang_text('result_analysis'), get_lang_text('report_export')]
        tabs = st.tabs(tab_titles)
        
        # 评估标签页
        with tabs[0]:
            try:
                for section, section_data in questionnaire.items():
                    with st.expander(get_display_section_title(section_data, st.session_state.language), expanded=True):
                        for q_id, question in section_data.get('questions', {}).items():
                            key = f"{section}_{q_id}"
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                type_class = {
                                    "PJ": "question-type-pj",
                                    "XO": "question-type-xo",
                                    "PW": "question-type-pw"
                                }.get(question["type"], "")
                                
                                description = get_translated_text(question.get('description', {}), st.session_state.language)
                                st.markdown(
                                    f'<span class="question-type {type_class}">{question["type"]}</span>'
                                    f'<span style="font-weight: bold;">{description}</span>',
                                    unsafe_allow_html=True
                                )
                            
                            with col2:
                                if question["type"] == "XO":
                                    # 是否题使用单选框
                                    current_value = st.session_state.responses.get(key, 0)
                                    yes_no_options = {
                                        'zh': {0: get_lang_text('answer_no'), 4: get_lang_text('answer_yes')},
                                        'en': {0: get_lang_text('answer_no'), 4: get_lang_text('answer_yes')}
                                    }
                                    st.session_state.responses[key] = st.radio(
                                        config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score'],
                                        options=[0, 4],
                                        format_func=lambda x: yes_no_options[st.session_state.language][x],
                                        horizontal=True,
                                        key=f"radio_{section}_{q_id}",
                                        label_visibility="collapsed",
                                        index=1 if current_value == 4 else 0
                                    )
                                elif question["type"] == "PJ":
                                    # 主观判断题使用下拉框
                                    current_value = st.session_state.responses.get(key, 0)
                                    score_labels = {
                                        'zh': {
                                            0: get_lang_text('not_implemented'),
                                            1: get_lang_text('preliminary_implementation'),
                                            2: get_lang_text('partial_implementation'),
                                            3: get_lang_text('mostly_implemented'),
                                            4: get_lang_text('fully_implemented')
                                        },
                                        'en': {
                                            0: get_lang_text('not_implemented'),
                                            1: get_lang_text('preliminary_implementation'),
                                            2: get_lang_text('partial_implementation'),
                                            3: get_lang_text('mostly_implemented'),
                                            4: get_lang_text('fully_implemented')
                                        }
                                    }
                                    # 修正index越界问题
                                    index = 0
                                    try:
                                        index = int(current_value)
                                        if index not in [0, 1, 2, 3, 4]:
                                            index = 0
                                    except Exception:
                                        index = 0
                                    st.session_state.responses[key] = st.selectbox(
                                        config.lang_zh['score'] if st.session_state.language == 'zh' else config.lang_en['score'],
                                        options=[0, 1, 2, 3, 4],
                                        format_func=lambda x: score_labels[st.session_state.language][x],
                                        key=f"select_{section}_{q_id}",
                                        label_visibility="collapsed",
                                        index=index  # 保证index合法
                                    )
                                else:  # PW类型
                                    # 多选题使用复选框
                                    if "sub_questions" in question:
                                        sub_questions = question.get('sub_questions', {}).get(st.session_state.language, [])
                                        for i, sub_q in enumerate(sub_questions, 1):
                                            sub_key = f"{key}_sub_{i}"
                                            if sub_key not in st.session_state.sub_responses:
                                                st.session_state.sub_responses[sub_key] = False
                                            checked = st.checkbox(
                                                sub_q,
                                                value=st.session_state.sub_responses.get(sub_key, False),
                                                key=f"checkbox_{section}_{q_id}_{i}_sub"
                                            )
                                            st.session_state.sub_responses[sub_key] = checked
                                        # 不再赋值st.session_state.responses[key]
                
                # 自动保存功能
                current_time = datetime.now()
                if (current_time - st.session_state.last_save_time).total_seconds() > 300:  # 每5分钟自动保存一次
                    try:
                        save_assessment_results(st.session_state.responses, st.session_state.sub_responses)
                        st.session_state.last_save_time = current_time
                        auto_save_text = config.lang_zh['progress_auto_saved'] if st.session_state.language == 'zh' else config.lang_en['progress_auto_saved']
                        st.toast(auto_save_text, icon="💾")
                    except Exception as e:
                        logging.error(f"自动保存失败: {str(e)}")
            
            except Exception as e:
                st.error(f"{get_lang_text('error_rendering_assessment')}: {str(e)}")
                logging.error(f"渲染评估页面失败: {str(e)}")
                logging.error(traceback.format_exc())

        # 结果分析标签页
        with tabs[1]:
            try:
                # 计算各部分得分
                section_scores = {}
                for section in questionnaire.keys():
                    section_responses = {k: v for k, v in st.session_state.responses.items() if k.startswith(section)}
                    section_sub_responses = {k: v for k, v in st.session_state.sub_responses.items() if k.startswith(section)}
                    
                    # 计算每个问题的得分
                    question_scores = []
                    questions = questionnaire[section].get('questions', {})
                    for q_id, question in questions.items():
                        key = f"{section}_{q_id}"
                        weight = score_weights_config['question_weights'][section].get(q_id, 1)
                        if question["type"] == "PW":
                            # 为PW类型题目准备子响应数据
                            pw_sub_responses = {}
                            for i in range(1, len(question.get("sub_questions", {}).get(st.session_state.language, [])) + 1):
                                sub_key = f"{key}_sub_{i}"
                                pw_sub_responses[sub_key] = section_sub_responses.get(sub_key, False)
                            score = calculate_compliance_score(
                                0,  # PW类型题目不需要主得分
                                question["type"],
                                pw_sub_responses,
                                score_weight=weight
                            )
                        elif key in section_responses:
                            score = calculate_compliance_score(
                                section_responses[key],
                                question["type"],
                                None,
                                score_weight=weight
                            )
                        else:
                            score = 0
                        question_scores.append(score)
                    
                    # 计算要素总分（由平均分改为总和）
                    section_scores[section] = sum(question_scores) if question_scores else 0
                
                # 显示总体合规分数
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    # 只显示1000分制得分
                    total_score = calculate_total_score(section_scores)
                    st.metric(
                        config.lang_zh['overall_score'] if st.session_state.language == 'zh' else config.lang_en['overall_score'],
                        f"{total_score:.1f}/1000"
                    )
                
                # 显示雷达图
                radar_chart = create_radar_chart(section_scores)
                if radar_chart:
                    st.plotly_chart(radar_chart, use_container_width=True)
                else:
                    st.warning(config.lang_en['cannot_generate_radar_chart'] if st.session_state.language == 'en' else config.lang_zh['cannot_generate_radar_chart'])
                # 显示详细得分（去除百分比进度条，仅保留分数）
                st.subheader(config.lang_zh['element_scores'] if st.session_state.language == 'zh' else config.lang_en['element_scores'])
                cols = st.columns(3)
                for i, (section, score) in enumerate(section_scores.items()):
                    with cols[i % 3]:
                        section_name = questionnaire[section]['name'][st.session_state.language]
                        st.metric(section_name, f"{score:.1f}")
            
            except Exception as e:
                st.error(f"{get_lang_text('error_rendering_results')}: {str(e)}")
                logging.error(f"渲染结果分析页面失败: {str(e)}")
                logging.error(traceback.format_exc())

        # 报告导出标签页
        with tabs[2]:
            try:
                # 第一行：PDF和DOCX报告按钮
                st.markdown("<br>", unsafe_allow_html=True)  # 增加顶部间距
                col1, col2 = st.columns(2)
                
                # PDF报告导出按钮
                with col1:
                    pdf_btn_text = config.lang_zh['generate_pdf_report'] if st.session_state.language == 'zh' else config.lang_en['generate_pdf_report']
                    if st.button(pdf_btn_text, key="generate_pdf_report"):
                        spinner_text = config.lang_zh['generating_pdf_report'] if st.session_state.language == 'zh' else config.lang_en['generating_pdf_report']
                        with st.spinner(spinner_text):
                            pdf_buffer = create_pdf_report(section_scores, questionnaire, st.session_state.responses, st.session_state.sub_responses)
                            if pdf_buffer:
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                filename = f"ISO55001_{'Assessment_Report' if st.session_state.language == 'en' else get_lang_text('report_title')}_{timestamp}.pdf"
                                download_label = config.lang_zh['download_pdf_report'] if st.session_state.language == 'zh' else config.lang_en['download_pdf_report']
                                st.download_button(
                                    label=download_label,
                                    data=pdf_buffer,
                                    file_name=filename,
                                    mime="application/pdf"
                                )
                                success_msg = config.lang_zh['pdf_report_generated'] if st.session_state.language == 'zh' else config.lang_en['pdf_report_generated']
                                st.success(success_msg)
                            else:
                                error_msg = config.lang_zh['failed_generating_pdf_report'] if st.session_state.language == 'zh' else config.lang_en['failed_generating_pdf_report']
                                st.error(error_msg)
                
                # DOCX报告导出按钮
                with col2:
                    docx_btn_text = config.lang_zh['generate_docx_report'] if st.session_state.language == 'zh' else config.lang_en['generate_docx_report']
                    if st.button(docx_btn_text, key="generate_docx_report"):
                        spinner_text = config.lang_zh['generating_docx_report'] if st.session_state.language == 'zh' else config.lang_en['generating_docx_report']
                        with st.spinner(spinner_text):
                            docx_buffer = create_doc_report(section_scores, questionnaire, st.session_state.responses, st.session_state.sub_responses)
                            if docx_buffer:
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                filename = f"ISO55001_{'Assessment_Report' if st.session_state.language == 'en' else get_lang_text('report_title')}_{timestamp}.docx"
                                download_label = config.lang_zh['download_docx_report'] if st.session_state.language == 'zh' else config.lang_en['download_docx_report']
                                st.download_button(
                                    label=download_label,
                                    data=docx_buffer,
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                success_msg = config.lang_zh['docx_report_generated'] if st.session_state.language == 'zh' else config.lang_en['docx_report_generated']
                                st.success(success_msg)
                            else:
                                error_msg = config.lang_zh['failed_generating_docx_report'] if st.session_state.language == 'zh' else config.lang_en['failed_generating_docx_report']
                                st.error(error_msg)
                
                # 第二行：Excel和Markdown报告按钮
                st.markdown("<br>", unsafe_allow_html=True)  # 行间距
                col3, col4 = st.columns(2)
                
                # Excel报告导出按钮
                with col3:
                    excel_btn_text = config.lang_zh['generate_excel_report'] if st.session_state.language == 'zh' else config.lang_en['generate_excel_report']
                    if st.button(excel_btn_text, key="generate_excel_report"):
                        spinner_text = config.lang_zh['generating_excel_report'] if st.session_state.language == 'zh' else config.lang_en['generating_excel_report']
                        with st.spinner(spinner_text):
                            excel_buffer = create_excel_report(section_scores, questionnaire, st.session_state.responses, st.session_state.sub_responses)
                            if excel_buffer:
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                filename = f"ISO55001_{'Assessment_Report' if st.session_state.language == 'en' else get_lang_text('report_title')}_{timestamp}.xlsx"
                                download_label = config.lang_zh['download_excel_report'] if st.session_state.language == 'zh' else config.lang_en['download_excel_report']
                                st.download_button(
                                    label=download_label,
                                    data=excel_buffer,
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                                success_msg = config.lang_zh['excel_report_generated'] if st.session_state.language == 'zh' else config.lang_en['excel_report_generated']
                                st.success(success_msg)
                            else:
                                error_msg = config.lang_zh['error_generating_excel_report'] if st.session_state.language == 'zh' else config.lang_en['error_generating_excel_report']
                                st.error(error_msg)
                
                # Markdown报告导出按钮
                with col4:
                    md_btn_text = config.lang_zh['generate_markdown_report'] if st.session_state.language == 'zh' else config.lang_en['generate_markdown_report']
                    if st.button(md_btn_text, key="generate_markdown_report"):
                        spinner_text = config.lang_zh['generating_markdown_report'] if st.session_state.language == 'zh' else config.lang_en['generating_markdown_report']
                        with st.spinner(spinner_text):
                            md_content = create_markdown_report(section_scores, questionnaire, st.session_state.responses, st.session_state.sub_responses)
                            if md_content:
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                filename = f"ISO55001_{'Assessment_Report' if st.session_state.language == 'en' else get_lang_text('report_title')}_{timestamp}.md"
                                download_label = config.lang_zh['download_markdown_report'] if st.session_state.language == 'zh' else config.lang_en['download_markdown_report']
                                st.download_button(
                                    label=download_label,
                                    data=md_content,
                                    file_name=filename,
                                    mime="text/markdown"
                                )
                                success_msg = config.lang_zh['markdown_report_generated'] if st.session_state.language == 'zh' else config.lang_en['markdown_report_generated']
                                st.success(success_msg)
                            else:
                                error_msg = config.lang_zh['failed_generating_markdown_report'] if st.session_state.language == 'zh' else config.lang_en['failed_generating_markdown_report']
                                st.error(error_msg)
            
            except Exception as e:
                st.error(f"{get_lang_text('error_generating_report')}: {str(e)}")
                logging.error(f"生成报告失败: {str(e)}")
                logging.error(traceback.format_exc())
    
    except Exception as e:
        st.error(f"应用运行出错: {str(e)}")
        logging.error(f"应用运行失败: {str(e)}")
        logging.error(traceback.format_exc())

if __name__ == "__main__":
    main() 